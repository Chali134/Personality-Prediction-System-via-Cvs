import os
import json
from fpdf import FPDF
from docx import Document

# Directory to save output files
output_dir = "output_files"
os.makedirs(output_dir, exist_ok=True)

# Full list of candidates with detailed data
candidates = {
    "Ali_Hassan": {
        "name": "Ali Hassan",
        "experience": 5,
        "skills": ["Python", "Machine Learning", "Data Analysis"],
        "soft_skills": ["Communication", "Problem-Solving"],
        "certifications": ["Certified Data Scientist"],
        "interests": ["Artificial Intelligence", "Gaming"],
        "personality_test_score": 85
    },
    "Zara_Malik": {
        "name": "Zara Malik",
        "experience": 3,
        "skills": ["JavaScript", "Web Development", "React"],
        "soft_skills": ["Adaptability", "Teamwork"],
        "certifications": ["Web Developer Certification"],
        "interests": ["Photography", "Travel"],
        "personality_test_score": 90
    },
    "Ayesha_Farooq": {
        "name": "Ayesha Farooq",
        "experience": 4,
        "skills": ["Java", "Spring Boot", "Microservices"],
        "soft_skills": ["Leadership", "Time Management"],
        "certifications": ["Java Certification"],
        "interests": ["Blogging", "Cycling"],
        "personality_test_score": 88
    },
    "Bilal_Khan": {
        "name": "Bilal Khan",
        "experience": 6,
        "skills": ["C++", "Embedded Systems"],
        "soft_skills": ["Critical Thinking", "Attention to Detail"],
        "certifications": ["Embedded Systems Certification"],
        "interests": ["Robotics", "Chess"],
        "personality_test_score": 92
    },
    "Umar_Shah": {
        "name": "Umar Shah",
        "experience": 2,
        "skills": ["HTML", "CSS", "UX/UI Design"],
        "soft_skills": ["Creativity", "Empathy"],
        "certifications": ["UX/UI Design Certification"],
        "interests": ["Design", "Music"],
        "personality_test_score": 87
    },
    "Sara_Ahmed": {
        "name": "Sara Ahmed",
        "experience": 4,
        "skills": ["SQL", "Data Warehousing"],
        "soft_skills": ["Collaboration", "Analytical Skills"],
        "certifications": ["Data Analyst Certification"],
        "interests": ["Cooking", "Traveling"],
        "personality_test_score": 89
    },
    "Kamran_Ali": {
        "name": "Kamran Ali",
        "experience": 5,
        "skills": ["Project Management", "Agile"],
        "soft_skills": ["Leadership", "Strategic Thinking"],
        "certifications": ["PMP Certification"],
        "interests": ["Reading", "Hiking"],
        "personality_test_score": 93
    }
}

# Function to save candidate data to text file
def save_to_text(name, candidate_data):
    try:
        with open(os.path.join(output_dir, f"{name}.txt"), "w") as file:
            file.write(json.dumps(candidate_data, indent=4))
    except Exception as e:
        print(f"Error saving text file for {name}: {e}")

# Function to save candidate data to PDF file with detailed formatting
def save_to_pdf(name, candidate_data):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=f"Candidate: {candidate_data['name']}", ln=True, align="L")
        
        for key, value in candidate_data.items():
            if isinstance(value, list):
                formatted_value = ", ".join(value)
            else:
                formatted_value = str(value)
            pdf.cell(200, 10, txt=f"{key.capitalize()}: {formatted_value}", ln=True)
            
        pdf.output(os.path.join(output_dir, f"{name}.pdf"))
    except Exception as e:
        print(f"Error saving PDF file for {name}: {e}")

# Function to save candidate data to Word file with detailed formatting
def save_to_word(name, candidate_data):
    try:
        doc = Document()
        doc.add_heading(f"Candidate: {candidate_data['name']}", level=1)
        for key, value in candidate_data.items():
            if isinstance(value, list):
                formatted_value = ", ".join(value)
            else:
                formatted_value = str(value)
            doc.add_paragraph(f"{key.capitalize()}: {formatted_value}")
        doc.save(os.path.join(output_dir, f"{name}.docx"))
    except Exception as e:
        print(f"Error saving Word file for {name}: {e}")

# Processing each candidate with assigned file format
for name, candidate in candidates.items():
    if name in ["Ali_Hassan", "Zara_Malik", "Ayesha_Farooq"]:  # TXT files
        save_to_text(name, candidate)
    elif name in ["Bilal_Khan", "Umar_Shah"]:  # PDF files
        save_to_pdf(name, candidate)
    elif name in ["Sara_Ahmed", "Kamran_Ali"]:  # Word files
        save_to_word(name, candidate)

# List files generated
print("Files generated:", os.listdir(output_dir))
